# commands/report.py
# تقرير وورد احترافي
# كل إعلان في جدول مستقل + صفحة مستقلة
# يولّد تقريرين: إيجار + بيع
# يستثني المؤجر (status == "rented")

import os
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional, Tuple, List

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import CommandHandler, CallbackQueryHandler, ContextTypes

from settings import load_config
from services.storage import load_data

from docx import Document
from docx.shared import Pt, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE


def add_hyperlink(paragraph, url: str, text: str, *, color_hex: str = "0000FF"):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    c = OxmlElement("w:color")
    c.set(qn("w:val"), color_hex)
    rPr.append(c)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _get_bot_username() -> str:
    cfg = load_config()
    return (os.getenv("BOT_USERNAME") or cfg.get("BOT_USERNAME") or "abo_alhassanbot").strip().lstrip("@")


def _tg_ad_link(ad_id: str):
    if not ad_id:
        return None
    bot_username = _get_bot_username()
    return f"https://t.me/{bot_username}?start=ad_{ad_id}"


def set_cell_shading(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def is_admin_user(user_id: int) -> bool:
    cfg = load_config()
    admin_id = cfg.get("ADMIN_ID")
    try:
        admin_id = int(admin_id)
    except (TypeError, ValueError):
        admin_id = None

    admins = cfg.get("ADMINS", [])
    admins = [int(a) for a in admins if str(a).isdigit()]
    return user_id == admin_id or user_id in admins


BASE_DIR = Path(__file__).resolve().parent.parent
REPORTS_DIR = BASE_DIR / "reports"
MEDIA_DIR = REPORTS_DIR / "media"
REPORTS_DIR.mkdir(parents=True, exist_ok=True)
MEDIA_DIR.mkdir(parents=True, exist_ok=True)

_ARABIC_DIGITS = str.maketrans("٠١٢٣٤٥٦٧٨٩", "0123456789")
_ARABIC_DIGITS_2 = str.maketrans("۰۱۲۳۴۵۶۷۸۹", "0123456789")


def _norm(text: str) -> str:
    if not text:
        return ""
    t = text.strip()
    t = t.translate(_ARABIC_DIGITS).translate(_ARABIC_DIGITS_2)
    t = t.replace("أ", "ا").replace("إ", "ا").replace("آ", "ا")
    return t


def _norm_lower(text: str) -> str:
    return _norm(text).lower()


def _first_non_empty_line(text: str) -> str:
    for line in (text or "").splitlines():
        s = line.strip()
        if s:
            return s
    return (text or "").strip()


def _short_title_from_text(text: str) -> str:
    if not text:
        return "إعلان"

    t = _norm(text)
    tl = t.lower()

    kind = ""
    if any(w in tl for w in ["شقه", "شقة", "شقق"]):
        kind = "شقة"
    elif any(w in tl for w in ["بيت", "دار", "منزل"]):
        if "طابقين" in tl:
            kind = "بيت طابقين"
        else:
            kind = "بيت"
    elif any(w in tl for w in ["ارض", "قطعة", "قطعه"]):
        kind = "قطعة أرض"
    elif any(w in tl for w in ["محل", "محلات"]):
        kind = "محل"
    elif any(w in tl for w in ["بناية", "عمارة"]):
        kind = "بناية"

    op = ""
    if any(w in tl for w in ["للبيع", "بيع"]):
        op = "بيع"
    elif any(w in tl for w in ["للايجار", "ايجار", "اجار"]):
        op = "إيجار"

    area = ""
    m = re.search(r"(?:في\s+منطقة|منطقة)\s+([^\n\r\-—]{2,30})", t)
    if m:
        area = m.group(1).strip()

    parts = []
    if kind:
        parts.append(kind)
    if op:
        parts.append(op)
    if area:
        parts.append(area)

    if not parts:
        line = _first_non_empty_line(t)
        if len(line) > 60:
            line = line[:60] + "..."
        return line

    return " - ".join(parts)


def _extract_price_iqd(text: str) -> Optional[int]:
    t = _norm_lower(text)
    clean = t.replace("٬", ",").replace("،", ",")
    clean = clean.replace(",", " ")
    clean = re.sub(r"\s+", " ", clean).strip()

    total = 0
    for m in re.finditer(r"([0-9]+(?:\.[0-9]+)?)\s*(مليار|مليون|الف|ألف)", clean):
        num = float(m.group(1))
        unit = m.group(2)
        if unit == "مليار":
            total += int(num * 1_000_000_000)
        elif unit == "مليون":
            total += int(num * 1_000_000)
        else:
            total += int(num * 1_000)

    word_base = 0
    if "مليار" in clean:
        if "مليارين" in clean or "ملياران" in clean:
            word_base += 2_000_000_000
        else:
            word_base += 1_000_000_000

    if "مليون" in clean:
        if "مليونين" in clean or "مليونان" in clean:
            word_base += 2_000_000
        else:
            word_base += 1_000_000

    if ("ونص" in clean or "ونصف" in clean):
        if "مليار" in clean:
            word_base += 500_000_000
        elif "مليون" in clean:
            word_base += 500_000
        elif ("الف" in clean or "ألف" in clean):
            word_base += 500

    total += word_base
    if total > 0:
        return int(total)

    m2 = re.search(r"([0-9]{4,})", clean)
    if m2:
        try:
            return int(float(m2.group(1)))
        except Exception:
            pass

    return None


def _format_iqd(n: Optional[int]) -> str:
    if n is None:
        return ""
    s = f"{n:,}"
    return s + "\u00A0" + "د.ع"


def _format_dt_cell(created_at: str) -> str:
    s = str(created_at or "").strip().replace("T", " ")
    dt = None

    try:
        dt = datetime.fromisoformat(str(created_at).replace("Z", "").replace("z", "").replace("T", " "))
    except Exception:
        dt = None

    if dt is None:
        for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M"):
            try:
                dt = datetime.strptime(s, fmt)
                break
            except Exception:
                continue

    if dt is None:
        parts = s.split()
        if len(parts) >= 2:
            return parts[0] + "\n" + parts[1]
        return s

    date_part = dt.strftime("%Y-%m-%d")
    time_part = dt.strftime("%I:%M %p").replace("AM", "ص").replace("PM", "م")
    if time_part.startswith("0"):
        time_part = time_part[1:]
    return f"{date_part}\n{time_part}"


def _media_summary(ad: dict) -> str:
    photos = ad.get("photos", []) or []
    videos = ad.get("videos", []) or []
    docs = ad.get("documents", []) or []
    return f"صور({len(photos)}) فيديو({len(videos)}) ملفات({len(docs)})"


async def _prepare_media_link(_bot, ad: dict) -> Tuple[str, Optional[str]]:
    photos = ad.get("photos", []) or []
    videos = ad.get("videos", []) or []
    docs = ad.get("documents", []) or []
    total = len(photos) + len(videos) + len(docs)
    if total == 0:
        return "لا يوجد", None

    summary = _media_summary(ad)
    ad_id = ad.get("ad_id") or uuid.uuid4().hex

    txt_path = MEDIA_DIR / f"{ad_id}.txt"
    lines = []
    if photos:
        lines.append(f"PHOTOS ({len(photos)}):")
        lines.extend(photos)
        lines.append("")
    if videos:
        lines.append(f"VIDEOS ({len(videos)}):")
        lines.extend(videos)
        lines.append("")
    if docs:
        lines.append(f"DOCUMENTS ({len(docs)}):")
        lines.extend(docs)
        lines.append("")

    txt_path.write_text("\n".join(lines).strip(), encoding="utf-8")
    return summary, str(txt_path)


def _collect_ads(data: dict, root_key: str) -> List[dict]:
    cats = data.get("categories", {})
    root = cats.get(root_key, {})
    out: List[dict] = []

    def walk(node):
        subs = node.get("sub", {})
        items = node.get("items", None)

        if isinstance(items, list) and not subs:
            for it in items:
                if it.get("status") == "rented":
                    continue
                out.append(it)
            return

        for child in subs.values():
            walk(child)

    walk(root)
    return out


def _add_report_header(doc: Document, title: str, business_name: str, now: datetime):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(business_name)
    run2.font.size = Pt(12)

    time_hdr = now.strftime("%Y-%m-%d   %I:%M %p").replace("AM", "ص").replace("PM", "م")
    if " 0" in time_hdr:
        time_hdr = time_hdr.replace(" 0", " ", 1)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(time_hdr)
    run3.font.size = Pt(11)

    doc.add_paragraph("")


def _add_single_ad_table(doc: Document, idx: int, ad: dict, media_summary: str):
    cols = 7
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"

    widths = [
        Inches(0.55),
        Inches(1.20),
        Inches(2.10),
        Inches(1.10),
        Inches(1.05),
        Inches(1.10),
        Inches(1.05),
    ]

    hdr = table.rows[0].cells
    headers = ["#", "التاريخ/\nالوقت", "العنوان", "السعر", "رقم الهاتف", "الموقع", "وسائط"]
    HEADER_FILL = "B7DEE8"

    for i, h in enumerate(headers):
        hdr[i].width = widths[i]
        hdr[i].text = h
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for para in hdr[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                r.bold = True
                r.font.size = Pt(10)
        set_cell_shading(hdr[i], HEADER_FILL)

    created_at = ad.get("created_at") or datetime.now().strftime("%Y-%m-%d %H:%M")
    text = ad.get("text", "")
    title_short = _short_title_from_text(text)
    price = _format_iqd(_extract_price_iqd(text))

    row = table.add_row().cells
    for i in range(cols):
        row[i].width = widths[i]
        row[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    row[0].text = str(idx)
    row[1].text = _format_dt_cell(str(created_at))
    row[2].text = title_short
    row[3].text = price
    row[4].text = str(ad.get("report_phone", "") or "")
    row[5].text = str(ad.get("report_location", "") or "")
    row[6].text = media_summary

    ad_id = ad.get("ad_id") or ""
    link = _tg_ad_link(ad_id)
    if link:
        pmp = row[6].add_paragraph()
        add_hyperlink(pmp, link, "عرض الإعلان في البوت", color_hex="0000FF")
        pmp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for c in [0, 1, 3, 4, 5, 6]:
        for para in row[c].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                r.font.size = Pt(9)

    for para in row[2].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in para.runs:
            r.font.size = Pt(9)

    row2 = table.add_row().cells
    merged = row2[0].merge(row2[-1])
    merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    ptxt = merged.paragraphs[0]
    ptxt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    runk = ptxt.add_run((text or "").strip())
    runk.font.size = Pt(10)


def _build_report_doc(*, title: str, business_name: str, ads: List[dict], media_links: List[Tuple[str, Optional[str]]]) -> Path:
    now = datetime.now()
    stamp = now.strftime("%Y-%m-%d_%H-%M-%S")
    filename = f"{title}_{stamp}.docx".replace(" ", "_")
    out_path = REPORTS_DIR / filename

    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(12.7)
    section.right_margin = Mm(12.7)
    section.top_margin = Mm(12.7)
    section.bottom_margin = Mm(12.7)

    if not ads:
        _add_report_header(doc, title, business_name, now)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("لا توجد إعلانات")
        r.font.size = Pt(12)
        doc.save(out_path)
        return out_path

    for idx, ad in enumerate(ads, start=1):
        _add_report_header(doc, title, business_name, now)
        media_summary, _ = media_links[idx - 1]
        _add_single_ad_table(doc, idx, ad, media_summary)

        if idx < len(ads):
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)

    doc.save(out_path)
    return out_path


async def report_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    if not is_admin_user(uid):
        await update.message.reply_text("❌ هذا الأمر للإدارة فقط.")
        return

    kb = InlineKeyboardMarkup([
        [InlineKeyboardButton("📄 تقرير الإيجارات", callback_data="report:rent")],
        [InlineKeyboardButton("📄 تقرير البيع", callback_data="report:sale")],
        [InlineKeyboardButton("❌ إلغاء", callback_data="report:cancel")],
    ])
    await update.message.reply_text("اختر نوع التقرير:", reply_markup=kb)


async def report_buttons(update: Update, context: ContextTypes.DEFAULT_TYPE):
    q = update.callback_query
    await q.answer()

    uid = q.from_user.id
    if not is_admin_user(uid):
        return

    if q.data == "report:cancel":
        try:
            await q.message.delete()
        except Exception:
            try:
                await q.message.edit_reply_markup(reply_markup=None)
            except Exception:
                pass
        return

    data = load_data()
    info = data.get("info", {})
    business_name = info.get("business_name", "عقارات")

    if q.data == "report:rent":
        ads = _collect_ads(data, "rent")
        title = "تقرير إيجار"
    elif q.data == "report:sale":
        ads = _collect_ads(data, "sale")
        title = "تقرير بيع"
    else:
        return

    media_links: List[Tuple[str, Optional[str]]] = []
    for ad in ads:
        summary, link = await _prepare_media_link(context.bot, ad)
        media_links.append((summary, link))

    try:
        path = _build_report_doc(title=title, business_name=business_name, ads=ads, media_links=media_links)
    except Exception as e:
        await q.message.reply_text(f"❌ خطأ إنشاء التقرير: {e}")
        return

    try:
        with open(path, "rb") as f:
            await q.message.reply_document(
                document=f,
                filename=path.name,
                caption=f"✅ {title} جاهز"
            )
    finally:
        try:
            os.remove(path)
        except Exception:
            pass


def register(app):
    app.add_handler(CommandHandler("report", report_cmd))
    app.add_handler(CallbackQueryHandler(report_buttons, pattern=r"^report:(rent|sale|cancel)$"))